import streamlit as st
import os
import requests
import docx
from datetime import datetime
import io

st.session_state.ended = False

# Custom CSS for styling
st.markdown("""
    <style>
    .title {
        font-size:30px !important;
        #color:#FF5733;
        color:#2E86C1;     
        text-align:center;
    }
    .subtitle {
        font-size:20px !important;
        color:#FF5733;
        text-align:left;
    }
    .highlight {
        color:#28B463;  
        font-weight:bold;
    }
    </style>
""", unsafe_allow_html=True)


st.markdown("<p class='title'>Risk Analysis Report Generator</p>", unsafe_allow_html=True)

sentence = """
This module generates a comprehensive Cyber Risk and Resiliency Report for industrial systems, using the results of a prior sensitivity analysis. 
 It prompts the user for an API key and business sector. The tool constructs a detailed prompt and sends it to an AI model to produce a report tailored for senior leadership. 
 The generated report is displayed in the tool and can be downloaded as a formatted Word document. 
 The report covers executive summary, risk analysis, sensitivity outcomes, and actionable recommendations, all based on the system’s dependency structure and risk metrics. 
"""

st.markdown(f"<p class='highlight'>{sentence}</p>", unsafe_allow_html=True)


if 'ended' not in st.session_state:
      st.session_state.ended = False

if st.session_state.ended:
      st.info("The Module produced the Reports. It is the last action to be performed. Please close this tab and return to the RISKED Menu tab.")
      st.stop()


# --- CONFIGURATION ---
DEFAULT_MODEL = "llama-3.1-8b-instant"
MODEL = DEFAULT_MODEL
API_KEY = "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"
required_files = [
    "one_point_sensitivity.csv",
    "data_model_with_posterior.csv"
]
# Hardcoded values
sector = "Manufacturing"

# current_datetime = datetime.now().strftime('%A, %B %d, %Y at %I:%M %p')

# --- Ask user for API key in sidebar ---
# st.sidebar.header("Groq/Llama API Key")
# api_key = st.sidebar.text_input(
#     "Enter your Groq/Llama API Key:",
#     type="password",
#     help="Paste your Groq/Llama API key here. It will be used for all API requests."
# )

# if not api_key:
#     st.warning("Please enter your Groq/Llama API key in the sidebar to continue.")
#     st.stop()


# In the sidebar, ask for company name
sector = st.sidebar.text_input(
    "Enter your organisation's Sector:",
    help="Type your organisation's Sector so as to focus the analysis."
)

if not sector:
    st.warning("Please enter your organisation's Sector in the sidebar to continue.")
    st.stop()


# PROMPT_TEMPLATE = """

# Business Sector: {sector}

# You are an AI expert in cybersecurity risk analysis for industrial systems (e.g., SCADA, ICS). Your expertise includes probabilistic modelling, sensitivity analysis, and interpreting how failures or attacks propagate through complex dependencies. You are familiar with standards such as MITRE ATT&CK for ICS, NIST SP 800-82, and ISA/IEC 62443.

# You are provided with two CSV files:
# - one_point_sensitivity.csv
# - data_model_with_posterior.csv

# **File Descriptions:**
# - *one_point_sensitivity.csv*: Results of single-node sensitivity analysis. The key column, `prob_given_node0`, shows the marginal probability of the root node if a given leaf node fails completely (zero percent reliability).
# - *data_model_with_posterior.csv*: Contains the Bayesian Network structure, node dependencies, and baseline probabilities, including the root node.

# **Task:**  
# Generate a comprehensive Cyber Risk and Resiliency Report for senior leadership (CISO, CRO, Asset Owners). The report should be accessible to non-technical readers but analytically rigorous. Use your own words and avoid copying phrases from this prompt.
# In producing your report, please mutiply all probabilities by 100 to convert them into percentages with two decimal places.
# **Report Structure:**

# 1. **Executive Summary**
#    - Briefly describe the system and its dependency model.
#    - Explain the meaning of the "root node" (e.g., overall system availability).
#    - State the current marginal probability of the root node and interpret its business impact.
#    - Summarise key system characteristics (dependency depth, complexity, coupling).
#    - Highlight the most critical leaf nodes and their impact.

# 2. **Probabilistic Risk Analysis**
#    - Identify 2 to 3 plausible failure scenarios based on sensitivity scores and dependencies.
#    - For each scenario, explain:
#      - What can go wrong (business operations impact).
#      - How likely it is (qualitative and quantitative assessment).
#      - Potential consequences (business, safety, regulatory, financial), referencing relevant frameworks.

# 3. **Sensitivity Analysis Outcomes**
#    - Use one_point_sensitivity.csv to identify the top 5 most critical leaf nodes (lowest Probability Sensitivity Scores).
#    - For each node, explain:
#      - How its failure affects the root node’s risk.
#      - How root node reliability improves if the node is fully reliable.
#      - Provide formatted Probability Sensitivity Scores (e.g., “Root node reliability improves to 12.34 percent if Node X is fully reliable”).
#    - Use data_model_with_posterior.csv to discuss interdependencies and cascading effects.

# 4. **Recommendations & Strategic Reminders**
#    - Suggest targeted mitigation strategies for the top 5 critical nodes, referencing standards where appropriate.
#    - Conclude with two paragraphs on essential security hygiene and foundational practices, referencing relevant frameworks.

# **Formatting Instructions:**
# - Format all probabilities and sensitivity scores as percentages with two decimal places (e.g., 21.34%).
# - Structure the report with clear headings and concise paragraphs.

# **Data Provided:**
# --- one_point_sensitivity.csv ---
# {one_point_sensitivity}
# --- data_model_with_posterior.csv ---
# {data_model_with_posterior}

# """

PROMPT_TEMPLATE = """

Business Sector: {sector}


You are a world-class expert in cybersecurity risk analysis for industrial control systems (ICS) and operational technology (OT) environments. You combine the expertise of a CISO, risk quantifier, and OT security engineer with 25+ years of experience. Your specialty is translating complex probabilistic models into strategic business insights for senior leadership.

## INPUT DATA CONTEXT

You are provided with two CSV files from a Bayesian Network analysis of an industrial control system:

### File 1: one_point_sensitivity.csv
This contains single-node sensitivity analysis results. The key column `prob_given_node0` shows the marginal probability of the root node (overall system reliability/availability) if a given leaf node fails completely (0% reliability). Lower values indicate higher criticality.

### File 2: data_model_with_posterior.csv
This contains the Bayesian Network structure, node dependencies, baseline probabilities, and the calculated posterior probability of the root node under current conditions.

[CSV DATA START]
--- one_point_sensitivity.csv ---
{one_point_sensitivity}
--- data_model_with_posterior.csv ---
{data_model_with_posterior}
[CSV DATA END]

## ANALYSIS METHODOLOGY

### Probability Formatting
Multiply all probabilities by 100 and format as percentages with two decimal places (e.g., 0.2134 → 21.34%).

### Criticality Scoring
Calculate a Criticality Score for each leaf node using:
```
Criticality Score = (Baseline Risk Contribution × Sensitivity Impact) / (Mitigation Complexity)
```
Where:
- **Baseline Risk Contribution**: Current probability contribution to root node failure
- **Sensitivity Impact**: Change in root node probability when node fails completely
- **Mitigation Complexity**: Estimated effort to secure the node (Low/Medium/High)

### Risk Categorization
Classify risks into:
- **🔴 Critical (15-20)**: Immediate threat to operations, requires action within 30 days
- **🟠 High (10-14)**: Significant vulnerability, address within quarter
- **🟡 Medium (5-9)**: Notable risk, plan for next planning cycle
- **🟢 Low (1-4)**: Acceptable risk or minor improvements

## REPORT STRUCTURE

### 1. EXECUTIVE SUMMARY

#### 1.1 System Overview
- **System Description**: Brief explanation of the industrial control system being analyzed
- **Dependency Model**: How components interconnect and depend on each other
- **Root Node Definition**: What the root node represents (e.g., "Overall System Availability" or "Critical Process Reliability")
- **Model Characteristics**: 
  - Number of nodes: [X]
  - Dependency depth: [Shallow/Moderate/Deep]
  - System coupling: [Loose/Moderate/Tight]

#### 1.2 Current Risk Posture
- **Baseline Root Node Reliability**: [XX.XX]%
- **Business Impact Interpretation**: What this reliability means for operations (e.g., "The system is expected to be fully operational XX% of the time, resulting in approximately YY hours of downtime annually")
- **Overall Risk Rating**: [🔴 Critical / 🟠 High / 🟡 Medium / 🟢 Low]
- **Maturity Level**: [Initial/Developing/Defined/Managed/Optimizing]

#### 1.3 Critical Findings at a Glance
| Risk Level | Count | Example Nodes |
|------------|-------|---------------|
| 🔴 Critical | X | [Node A, Node B] |
| 🟠 High | Y | [Node C, Node D] |
| 🟡 Medium | Z | [Node E, Node F] |
| 🟢 Low | W | [Node G, Node H] |

#### 1.4 Key Insights
- **Greatest Strength**: [Area with lowest risk contribution]
- **Biggest Vulnerability**: [Most critical node with explanation]
- **Quick Win Opportunity**: [Node where small investment yields significant improvement]
- **Cascading Risk Concern**: [Node whose failure triggers multiple downstream effects]

---

### 2. DETAILED RISK ANALYSIS

#### 2.1 Top 5 Most Critical Nodes

For each of the top 5 most critical nodes (based on sensitivity analysis and dependencies):

**Node [Name]**
- **Risk Rating**: [🔴/🟠/🟡/🟢]
- **Criticality Score**: [XX/20]
- **Current State**: [Description of node's role and current reliability]
- **Failure Impact**: 
  - Root node reliability if this node fails: [XX.XX]%
  - Degradation from baseline: [ -XX.XX]%
  - Business operations impact: [Specific operational consequences]
- **Cascading Effects**: [Which downstream nodes are affected]
- **Attack Vectors**: [How this node could be compromised, referencing MITRE ATT&CK for ICS]
- **Detection Difficulty**: [Easy/Moderate/Hard - how quickly failures would be detected]

#### 2.2 Scenario Analysis

Identify and analyze 3 plausible failure scenarios:

**Scenario: [Descriptive Name]**
- **Triggering Node(s)**: [Primary node(s) involved]
- **Attack Path**: [Step-by-step failure/attack propagation]
- **Likelihood**: [Qualitative + Quantitative assessment]
- **Time to Critical Impact**: [How quickly operations would be affected]
- **Business Consequences**:
  - Operational: [Production impact, downtime estimate]
- **Worst-case Outcome**: [Most severe potential result]



#### 2.3 Dependency Chain Analysis

- **Most Vulnerable Chain**: [Path through system with highest cumulative risk]
- **Single Point of Failure**: [Nodes whose failure would collapse multiple chains]
- **Redundancy Assessment**: [Where redundancy exists vs. where it's needed]
- **Resilience Score**: [Assessment of system's ability to absorb failures]

---

### 3. SENSITIVITY ANALYSIS OUTCOMES

#### 3.1 Critical Node Ranking

```
Rank | Node Name | Current Root | If Node Fails | Delta | Criticality
     |           | Reliability  | Reliability   |       | Score
-----|-----------|--------------|---------------|-------|------------
1    | [Node]    | [XX.XX]%     | [XX.XX]%      | [X.XX] | [XX]/20
2    | [Node]    | [XX.XX]%     | [XX.XX]%      | [X.XX] | [XX]/20
3    | [Node]    | [XX.XX]%     | [XX.XX]%      | [X.XX] | [XX]/20
4    | [Node]    | [XX.XX]%     | [XX.XX]%      | [X.XX] | [XX]/20
5    | [Node]    | [XX.XX]%     | [XX.XX]%      | [X.XX] | [XX]/20
```

# #### 3.2 Sensitivity Heat Map (Conceptual)

# ```
# Node Criticality Distribution:
                    
#     High Criticality          Low Criticality
#     ┌────────────────────────────────────┐
# Node1│████████████████████░░░░░░░░░░░░░░│ 85 percentile
# Node2│██████████████░░░░░░░░░░░░░░░░░░░░│ 72 percentile
# Node3│██████████████████████████████░░░░│ 94 percentile
# Node4│██████████░░░░░░░░░░░░░░░░░░░░░░░░│ 48 percentile
# Node5│████████████████████████░░░░░░░░░░│ 78 percentile
#     └────────────────────────────────────┘
#     0                                 100
#     (Percentile rank of criticality)
# ```

#### 3.2 Key Sensitivity Insights

- **Most Sensitive Node**: [Node] - Root reliability drops by [XX.XX]% when this node fails
- **Most Resilient Node**: [Node] - Minimal impact on root when failed
- **Unexpected Dependency**: [Surprising relationship revealed by analysis]
- **Mitigation Leverage Point**: [Node where improvement yields highest ROI]

---

### 4. STRATEGIC RECOMMENDATIONS

#### 4.1 Immediate Actions

**Priority 1: [Action for most critical node]**
- **Target Node**: [Node name]
- **Rationale**: [Why this action first - based on criticality score]
- **Recommended Controls**: 
  - [Specific control, referencing framework: e.g., "Implement network segmentation per ISA/IEC 62443-3-2"]
  - [Specific control]
- **Expected Improvement**: Root reliability increases to [XX.XX]% (gain of [X.XX]%)
- **Implementation Effort**: [⚡ Low / ⚡⚡ Medium / ⚡⚡⚡ High]
- **Success Metric**: [How to measure effectiveness]

---

### 5. CONCLUSION & CALL TO ACTION

#### 5.1 Current State Summary
The analysis reveals a [describe overall posture - e.g., "moderately resilient but critically vulnerable in specific areas"] industrial control system. With baseline root reliability of [XX.XX]%, the organization faces [describe business impact - e.g., "approximately Y hours of unexpected downtime annually, costing an estimated $Z in lost production"].

#### 5.2 Target State
By implementing the recommended actions, the organization can achieve:
- **Improved Root Reliability**: [XX.XX]% (gain of [X.XX]%)
- **Reduced Annual Expected Loss**: From $[X] to $[Y]
- **Enhanced Regulatory Compliance**: [Specific frameworks achieved]
- **Strengthened Competitive Position**: [Business advantage]



#### 5.3 Immediate Next Steps
1. **Approve funding** for Priority 1 actions ([estimated cost])
2. **Assign ownership** to [role] for critical node remediation
3. **Schedule deep-dive** on [specific area] within 30 days
4. **Present this report** to [board/risk committee/leadership team]

---

## REPORT FORMATTING REQUIREMENTS

- Use professional executive report language - clear, concise, and impactful
- Highlight critical findings with **bold** or 🔴 icons
- Use bullet points for readability and scanning
- Reference specific frameworks and standards in [brackets]
- Include effort estimates (⚡ Low, ⚡⚡ Medium, ⚡⚡⚡ High)
- Add risk icons (🔴 Critical, 🟠 High, 🟡 Medium, 🟢 Low)
- Include confidence levels for estimates where appropriate (High/Medium/Low)

## QUALITY CHECKS

Before finalizing, verify that your report:
1. Translates technical probabilities into business impact
2. Prioritizes recommendations by risk reduction potential
3. Includes actionable, specific next steps
4. References relevant frameworks appropriately
5. Provides both quantitative and qualitative analysis
6. Is accessible to non-technical senior leaders
7. Includes clear calls to action with ownership suggestions
8. Quantifies the cost of inaction
9. Provides ROI justification for investments
10. Maintains professional, executive-level tone

"""

def read_file_as_text(filename):
    try:
        with open(filename, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return None
    
api_key = "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"
def call_groq(prompt, api_key, model):
    url = "https://api.groq.com/openai/v1/chat/completions"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an expert in cybersecurity risk analysis for complex industrial systems like SCADA and ICS. You combine the skill and knowledge of over 100 top-tier cybersecurity Risk experts. "
                    "Your specialty is using Probability Risk Analysis, Bayesian Theory of Inference (Variable Elimination), Sensitivity Analysis & Parallelisation, probabilistic models to understand how failures and attacks can propagate through interdependent components. "
                    "Provide detailed, actionable analysis and recommendations. Be comprehensive and professional. Let your output be professionally formatted for the audience of senior security leadership (CISO, CRO, Asset Owners)."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.6,
        "max_tokens": 4000
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=120)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"API Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Unexpected Error: {str(e)}"

def generate_word_report(report, sector):
    doc = docx.Document()
    now = datetime.now().strftime('%A, %B %d, %Y at %I:%M %p')
    doc.add_heading('Executive Risk & Resiliency Report', 0)
    doc.add_paragraph(f"Business Sector: {sector}")
    doc.add_paragraph(f"Generated on: {now}")
    doc.add_paragraph("")
    lines = report.split('\n')
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == "---":
            doc.add_paragraph("")
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip()[2:-2]).bold = True
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("Generate and Download Report"):
    # Step 1: Check files
    missing_files = [f for f in required_files if not os.path.isfile(f)]
    
    if missing_files:
        st.error(f"The required file(s) to perform this process are missing."
        "You may have missed a step in the analysis sequence. "
        "Please complete the previous analysis steps to generate these files before proceeding.")
        st.stop()
    
    else:
        file_texts = {filename: read_file_as_text(filename) or "" for filename in required_files}
        prompt = PROMPT_TEMPLATE.format(
            sector=sector,
            one_point_sensitivity=file_texts["one_point_sensitivity.csv"],
            data_model_with_posterior=file_texts["data_model_with_posterior.csv"]
        )
        report = call_groq(prompt, API_KEY, MODEL)
        if report.startswith("API Error") or report.startswith("Unexpected Error"):
            st.error(f"Please provide a valid API key to proceed. You may go to Groq/Llama to generate a free API Key, if you do not have one.")
            st.stop()
            #st.error(report)
        else:
            # Display the report on the screen
            st.subheader("Generated Risk Analysis Report")
            st.markdown(report)
            # Provide download button for Word doc
            word_buffer = generate_word_report(report, sector)
            st.download_button(
                label="Download Risk Analysis Report (Word)",
                data=word_buffer,
                file_name=f"Risk_Resiliency_Report_{sector.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
st.caption("This is a PoC and it is powered by Groq and Llama 3. In the full version, this is powered by our purpose-built genAI for Cyber risk identification and assessement.")
st.markdown("---")    
if st.button("End Module"):
      st.session_state.ended = True
      st.rerun()
      