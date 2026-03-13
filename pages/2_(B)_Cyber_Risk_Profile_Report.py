import streamlit as st
import pandas as pd
import requests
from datetime import datetime
import docx
import io
import os

st.session_state.ended = False
# if 'ended' not in st.session_state:
#       st.session_state.ended = False

# if st.session_state.ended:
#       st.info("The Cyber Risk Security Profile Report Download is completed.")
#       st.stop()


# Custom CSS for styling
st.markdown("""
    <style>
    .title {
        font-size:35px !important;
        #color:#FF5733;
        color:#2E86C1;     
        text-align:left;
    }
    .subtitle {
        font-size:30px !important;
        color:#FF5733;
        text-align:centre;
    }
    .highlight {
        font-size:20px !important;
        color:#28B463;  
        font-weight:bold;
    }
    </style>
""", unsafe_allow_html=True)
st.session_state.ended = False

st.markdown("<p class='title'>Cyber Risk Profile Report Generator</p>", unsafe_allow_html=True)


# ========================
# CONFIGURATION
# ========================

COMPANY_NAME = "Demo - Cyber Risk Analytical Methods (CRAM) Inc"
AI_PROVIDER = "groq"
MODEL = "llama-3.1-8b-instant"
CSV_FILE = "Security_Profile_Responses.csv"
API_KEY = "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"

# --- NEW: Ask user for API key in sidebar ---

  
# st.sidebar.header("Groq/Llama API Key")
# API_KEY = st.sidebar.text_input(
#     "Enter your Groq/Llama API Key:",
#     type="password",
#     help="Paste your Groq/Llama API key here. It will be used for all API requests."
# )

# if not API_KEY:
#     st.warning("Please enter your Groq/Llama API key in the sidebar to continue.")
#     st.stop()


# In the sidebar, ask for company name
COMPANY_NAME = st.sidebar.text_input(
    "Enter your Company Name:",
    help="Type your organisation's name as you want it to appear in the report."
)

if not COMPANY_NAME:
    st.warning("Please enter your company name in the sidebar to continue.")
    st.stop()



# ========================
# API CALL
# ========================
def call_groq(prompt, api_key, model):
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a very senior cybersecurity expert specializing in cyber risk in operational technology environments and ICS/OT security frameworks and risk assessment. Provide detailed, actionable analysis and recommendations. Be comprehensive, exhaustive and professional in your analysis."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.5,
        "max_tokens": 4000
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=120)
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            return f"❌ API Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"❌ Unexpected Error: {str(e)}"

def generate_cybersecurity_analysis(csv_content, api_key, model):
    prompt = f"""
You are a world-class cybersecurity risk and compliance expert with 20+ years of experience in OT/ICS security. You have deep expertise implementing and auditing against the following 12 frameworks: 
ISA/IEC 62443 Series, NIST SP 800-82, ISO/IEC 27001, NIST Cybersecurity Framework (CSF), Cyber Essentials Plus, CIS Critical Security Controls, ENISA ICS Security Guidelines, CPNI Guidance, NIS2 Directive & EU Cyber Resilience Act, ISO/IEC 27019, NIST SP 800-53, ISA/IEC 61511.

## INPUT DATA CONTEXT
Below is a cybersecurity profile assessment containing security statements, categories, weights (1-4, where 4=highest criticality), user responses, and evidence:

[CSV DATA START]
{csv_content}
[CSV DATA END]

## SCORING METHODOLOGY
Calculate the Overall Cybersecurity Maturity Score using this weighted formula:
- Strongly Agree = 100% compliant (use evidence to verify)
- Agree = 75% compliant (use evidence to verify)
- Neutral = 50% compliant
- Disagree = 25% compliant
- Strongly Disagree = 0% compliant
- Not Applicable = Exclude from calculation

Weighted Score = Σ(Compliance % × Weight) / Σ(Weight) × 100

## EVIDENCE ASSESSMENT CRITERIA
Evaluate evidence quality for "Strongly Agree" and "Agree" responses:
- **Strong Evidence**: Specific documents, policies, technical implementations cited
- **Moderate Evidence**: General references without specific details
- **Weak Evidence**: Claims without supporting documentation
Flag "Strongly Agree" responses with weak/no evidence as potential overstatement.

## RISK RATING SCALE
Assign risk levels based on weighted gaps:
- **Critical (15-20)**: Immediate board-level attention, active threat potential
- **High (10-14)**: Significant vulnerability requiring quarterly remediation
- **Medium (5-9)**: Non-critical gaps, plan for next planning cycle
- **Low (1-4)**: Minor improvements, best practice enhancements

## COMPARATIVE ANALYSIS FRAMEWORK

### Industry Benchmark Data (For Reference)
Use these industry average benchmarks for comparison:
- **Critical Infrastructure/OT Sector Average**: 67%
- **Manufacturing Industry Average**: 62%
- **Energy/Utilities Sector Average**: 71%
- **Similar-sized Organizations (<500 employees)**: 58%
- **Similar-sized Organizations (500-5000 employees)**: 65%
- **Similar-sized Organizations (>5000 employees)**: 73%

### Peer Group Comparison Metrics
Compare against:
1. **Industry Sector Peers**: Organizations in same vertical
2. **Size-based Peers**: Organizations with similar employee count/revenue
3. **Framework-specific Benchmarks**: Average compliance by framework
4. **Regional/Regulatory Peers**: Organizations under similar regulations

### Maturity Model Levels
Position the organization on this maturity continuum:
- **Level 1: Initial** (0-25%) - Ad-hoc, reactive
- **Level 2: Developing** (26-50%) - Repeatable but informal
- **Level 3: Defined** (51-70%) - Standardized processes
- **Level 4: Managed** (71-85%) - Measured and monitored
- **Level 5: Optimizing** (86-100%) - Continuously improving

## ANALYSIS REQUIREMENTS

### 1. Executive Summary
- Brief overview of cybersecurity posture
- Key strengths with framework references
- Critical weaknesses with business impact
- **Overall Cybersecurity Maturity Score**: [Calculate using methodology above]
- **Maturity Level**: [Level 1-5 with description]
- **Peer Comparison**: [Above/Below/At industry average by X%]
- **Risk Rating Summary**: [Number of Critical/High/Medium/Low findings]

### 2. Performance Dashboard

#### 2.1 Framework Compliance Comparison

```
Metric                    | Your Score | Industry Avg | Variance | Trend
--------------------------|------------|--------------|----------|--------
Overall Maturity          | [XX]%      | [XX]%        | +/-XX%   | 📈/📉/📊
System & Resilience       | [XX]%      | [XX]%        | +/-XX%   | 📈/📉/📊
Technology Controls       | [XX]%      | [XX]%        | +/-XX%   | 📈/📉/📊
Process & Governance      | [XX]%      | [XX]%        | +/-XX%   | 📈/📉/📊
People & Awareness        | [XX]%      | [XX]%        | +/-XX%   | 📈/📉/📊
```




Compare compliance levels against industry averages for each framework:
- ISA/IEC 62443 Series: [Your %] vs Industry [XX]%
- NIST SP 800-82: [Your %] vs Industry [XX]%
- ISO/IEC 27001: [Your %] vs Industry [XX]%
- NIST CSF: [Your %] vs Industry [XX]%
- Cyber Essentials Plus: [Your %] vs Industry [XX]%
- CIS Controls: [Your %] vs Industry [XX]%
- ENISA Guidelines: [Your %] vs Industry [XX]%
- NIS2 Compliance: [Your %] vs Industry [XX]%



#### 2.2 Key Insights from Visual Analysis
- **Your Position**: [Brief description of where you sit in the matrix]
- **Biggest Gap**: [Area with largest variance from industry average]
- **Greatest Strength**: [Area where you exceed peers]
- **Quick Win Opportunity**: [Area where small investment yields big improvement]



### 3. Detailed Gap Analysis & Improvement Plan

For each section below, include:
- **Risk Rating**: [Critical/High/Medium/Low]
- **Weighted Gap Score**: [Sum of weights for non-compliant items]
- **Framework References**: [Specific standards violated]
- **Peer Comparison**: [How peers typically perform in this area]

#### A. System & Organisational Resilience
- **Critical Gaps**: [Prioritize by weight, include specific framework mappings]
- **Evidence Quality Issues**: [Flag any overstated compliance]
- **Peer Context**: [Industry norm for these controls, what leaders do differently]
- **Recommendations**: [Prioritized actions with effort estimates (⚡ Low, ⚡⚡ Medium, ⚡⚡⚡ High)]

#### B. Technology & Technical Controls
- **Critical Gaps**: [Include missing technical controls, vulnerabilities]
- **Evidence Quality Issues**: [Verify technical implementations claimed]
- **Peer Context**: [Technology adoption rates among peers, emerging trends]
- **Recommendations**: [Specific technical improvements with priorities]

#### C. Process & Governance
- **Critical Gaps**: [Policy/procedure deficiencies, governance weaknesses]
- **Evidence Quality Issues**: [Review documentation quality]
- **Peer Context**: [Governance maturity of similar organizations]
- **Recommendations**: [Process improvements with timeline estimates]

#### D. People & Awareness
- **Critical Gaps**: [Training deficiencies, awareness gaps]
- **Evidence Quality Issues**: [Verify training records, awareness metrics]
- **Peer Context**: [Industry benchmark for training frequency and coverage]
- **Recommendations**: [Culture and capability building initiatives]

### 4. Regulatory & Compliance Impact
- Map gaps to specific regulatory requirements (NIS2, GDPR, etc.)
- Identify potential compliance violations and consequences
- Estimate compliance score for each major framework
- **Regulatory Risk vs Peers**: [How your compliance posture compares]

### 5. Strategic Roadmap with Competitive Context

**Immediate Priorities (Next 3–6 Months) - Quick Wins to Close Peer Gaps:**
1. [Action with risk reduction estimate, e.g., "Reduce critical risk by 40%"]
2. [Action that addresses biggest gap vs peers]
3. [Action with dependencies identified]
   *Expected outcome*: [Move from Xth to Yth percentile among peers]

**Short-term Goals (6–12 Months) - Match Industry Leaders:**
1. [Goal with milestone definition]
2. [Goal with success metrics]
3. [Goal to achieve parity with top quartile performers]
   *Expected outcome*: [Achieve industry-leading position in specific domains]

**Long-term Strategic Goals (12–24 Months) - Establish Competitive Advantage:**
1. [Strategic initiative with business case]
2. [Transformational goal with ROI estimate]
3. [Innovation areas that exceed regulatory requirements]
   *Expected outcome*: [Become recognized as industry leader]

### 6. Risk Quantification with Market Context
- **Current Risk Exposure**: [Estimate using weighted gaps]
- **Residual Risk After Recommendations**: [Projected improvement]
- **Top 5 Risks**: [Ranked by weight and business impact]
- **Market Risk Comparison**: [How your risk profile compares to competitors]
- **Cyber Insurance Implications**: [Expected premium impact vs market rates]

# ### 7. Investment Rationale & Business Case

# #### 7.1 Cost of Inaction
# - Potential breach costs based on industry data
# - Regulatory fine exposure compared to peers
# - Reputational damage relative to competitors
# - Customer/partner confidence impact

# #### 7.2 ROI Projections
# - Risk reduction ROI: [Expected % risk reduction per $ invested]
# - Competitive advantage value: [Market differentiation benefit]
# - Operational efficiency gains: [Process improvements vs peers]
# - Insurance premium reductions: [Expected savings]

# #### 7.3 Peer Success Stories
# - Case Study 1: [Similar organization that improved from X% to Y%]
# - Case Study 2: [Peer that avoided breach through similar investments]
# - Industry trend: [What top performers are investing in]

### 8. Conclusion & Call to Action
- Summarise the overall risk posture and the expected impact of following the recommendations.
- Highlight how closing the identified gaps will improve compliance, resilience, and reduce risk exposure.
- **Current Position**: Summary of risk posture and competitive standing
- **Target Position**: Where you could be in 12-24 months
- **The Gap**: What separates you from industry leaders
- **Business Case**: Why act now (competitive pressure, regulatory deadlines, threat landscape)
- **Expected Outcomes**: Following recommendations will:
  - Move you from [Current percentile] to [Target percentile] among peers
  - Reduce risk exposure by [X]%
  - Achieve compliance with [Y] frameworks
  - Position you as industry leader in [Z] areas
- **Next Steps**: Immediate actions for management

## FORMATTING REQUIREMENTS
- Use professional executive report style
- Include data tables and comparison matrices where helpful
- Highlight critical findings with **bold**
- Use bullet points for readability
- Include framework codes in [brackets] after relevant findings
- Add effort estimates (🔴 High, 🟡 Medium, 🟢 Low)
- Add risk icons (🔴 Critical, 🟠 High, 🟡 Medium, 🟢 Low)
- Add comparison indicators (📈 Above average, 📉 Below average, 📊 At average)

## QUALITY CHECKS
Before finalizing, verify:
1. All weights are properly considered in prioritization
2. Evidence supports claimed compliance levels
3. Recommendations address root causes, not symptoms
4. Framework mappings are accurate and specific
5. Business impact is clearly articulated
6. Peer comparisons are realistic and contextual
7. Competitive positioning is actionable
8. Investment case is compelling for executives
"""
    return call_groq(prompt, api_key, model)

def create_word_document(analysis_text, response_df):
    doc = docx.Document()
    doc.add_heading('DEMO - Cyber Risk Profile Analysis', 0)
    doc.add_paragraph(f"Organization: {COMPANY_NAME}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    lines = analysis_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == '---':
            doc.add_paragraph()
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.strip()[2:-2]).bold = True
        elif line.strip():
            doc.add_paragraph(line)
        i += 1
    doc.add_paragraph()
    doc.add_heading('Assessment Methodology', level=2)
    doc.add_paragraph("This analysis is based on 12 industry-standard cybersecurity frameworks and standards:")
    frameworks = [
        "ISA/IEC 62443 Series", "NIST SP 800-82", "ISO/IEC 27001", "NIST Cybersecurity Framework (CSF)",
        "Cyber Essentials Plus", "CIS Critical Security Controls (CIS Controls)", "ENISA ICS Security Guidelines",
        "CPNI Guidance", "NIS2 Directive & EU Cyber Resilience Act", "ISO/IEC 27019", "NIST SP 800-53", "ISA/IEC 61511"
    ]
    for framework in frameworks:
        doc.add_paragraph(framework, style='List Bullet')
    doc.add_paragraph()
    doc.add_heading('Raw Response Data', level=2)
    table = doc.add_table(rows=1, cols=len(response_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(response_df.columns):
        hdr_cells[i].text = str(column)
    for index, row in response_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    return doc

# ========================
# STREAMLIT APP
# ========================

#st.markdown("<p class='title'>Cybersecurity Risk Profile Report Generator</p>", unsafe_allow_html=True)
st.markdown("<p class='subtitle'>Cyber Risk Profile Report</p>", unsafe_allow_html=True)

# st.set_page_config(page_title="Cyber-risk Security Profile Report", layout="wide")
# st.title("Cybersecurity Profile Report")
coyx = (f"Organization:  {COMPANY_NAME}")
st.markdown(f"<p class='highlight'>{coyx}</p>", unsafe_allow_html=True)

# Automatically process the report
if os.path.exists(CSV_FILE):
    response_df = pd.read_csv(CSV_FILE)
    required_columns = ['Category', 'Security Profile Statement', 'Weight', 'Response', 'Evidence']
    missing_columns = [col for col in required_columns if col not in response_df.columns]
    if missing_columns:
        st.error(f"❌ Error: Missing required columns in CSV: {missing_columns}")
    else:
        csv_content = response_df.to_csv(index=False)
        analysis_result = generate_cybersecurity_analysis(csv_content, API_KEY, MODEL)
           
        
        if analysis_result.startswith("❌ API Error: 401") or "invalid_api_key" in analysis_result.lower():
            st.error(
                "Please provide a valid API key to proceed. "
                "You may go to Groq/Llama to generate a free API Key, if you do not have one."
                )
        elif analysis_result.startswith("❌"):
            st.error(analysis_result)
   
            
        else:
            #st.header("Cybersecurity Analysis Report")
            st.markdown(analysis_result)
            # Generate Word document
            doc = create_word_document(analysis_result, response_df)
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')
            safe_company_name = COMPANY_NAME.replace(' ', '_').replace('/', '_')
            output_filename = f"cybersecurity_analysis_report_{safe_company_name}_{timestamp}.docx"
            st.download_button(
                label="Download Report (Word doc)",
                data=doc_io,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.error(f"The required file to generate this report was not found. You may have missed a step in the analysis sequence. Please ensure you have performed previous steps.")

st.caption("This is a PoC and it is powered by Groq and Llama 3. In the full version, this is powered by our purpose-built genAI for Cyber risk identification and assessement (CRAMLLM).")

st.markdown("---")    
if st.button("End Module"):
      st.session_state.ended = True
      st.info("The Cyber Risk Security Profile Report Download is completed.")
      #st.rerun()